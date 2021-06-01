# -*- coding: utf-8 -*-
"""
Created on Tue Jun  1 14:04:29 2021

@author: Say My Name
"""
import numpy as np
import json
from docx import Document

def creating_list(arr,key):
    document=Document()
    document.add_heading('The LIST',0)
    for i,val in enumerate(arr):
        p = document.add_paragraph(str(key)+" no. "+str(i+1), style='List Bullet')
        for lis in val:
            #key1, value = list(lis.items())[0]
            print(lis,val[lis])
            p = document.add_paragraph(str(lis)+" : "+str(val[lis]), style='ListBullet2')
    
    document.save('list2.docx')
f = open('sample4.json',)
# returns JSON object as a dictionary
data = json.load(f)
  
  
# Closing file
f.close()  
#converting json dic to array
key, arr = list(data.items())[0]

creating_list(arr,key)